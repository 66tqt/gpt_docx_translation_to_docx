# 导入需要的库
from docx import Document
from docx.shared import Pt
import openai
import threading
import time






# 获取用户输入的API key
api_key = input("请输入API key: ")

# 尝试获取用户输入的base URL，如果没有则使用默认值
base_url_input = input("请输入base URL (如果没有，请直接回车): ")

# 根据用户输入创建client实例
if base_url_input:  # 如果用户输入了base URL
    client = openai.OpenAI(api_key=api_key, base_url=base_url_input)
else:  # 如果用户没有输入base URL，使用默认的构造函数
    client = openai.OpenAI(api_key=api_key)

# 获取用户输入的线程数
xian_cheng_int = int(input("请输入线程数（线程数越多，单位时间内翻译的段落越多）（若出现故障/api限流问题，优先使用1个线程）: "))

print("客户端配置完成。")














def get(j,number_paragraph):
    global my_json_data
    global doc
    global lock

    # 获取第一段的内容
    one_paragraph_text = doc.paragraphs[number_paragraph].text
    print(one_paragraph_text)

    # 创建一个对话列表
    message = [
        {"role": "system", "content": f"""Translate the following English (or other languages) into Chinese.If there is no text to translate, please copy the user's input as a reply.text:"{one_paragraph_text} " """}
    ]



    try:
        # 使用GPT-3.5对话模型进行翻译
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=message
        )
        # print(response)
        print(response.choices[0].message.content)

        with lock:
            # 获取第一段的第一个字符的字体大小
            one_paragraph = doc.paragraphs[number_paragraph]
            one_run = one_paragraph.runs[0]  # 获取第一个run,也就是0
            font_size = one_run.font.size  # 获取字体大小

            # 清除第一段的内容
            one_paragraph.clear()

            # 修改内容
            new_content = one_paragraph.add_run(response.choices[0].message.content)  # 添加新的内容

            # 设置字体大小
            if font_size:  # 判断是否获取到字体大小
                new_content.font.size = font_size
            else:
                new_content.font.size = Pt(11)  # 设置为默认的11磅字体大小

            # 保存修改后的文件
            doc.save('Translation_file.docx')  # modified_example.docx是保存的文件名，可以根据需要进行修改
            my_json_data['ingT'].remove(j)
            print("已保存文件")
    except:
        print("API调用超时或其他错误")

### 主函数

# 创建全局锁
lock = threading.Lock()


file = input("当前文件夹的docx文件名（无后缀）：")
print(file)

# 打开已有的docx文件
doc = Document(f'{file}.docx')  # example.docx是要修改的文件名，根据实际情况修改

json_file = f"{file}.json"  # 设置要检测的文件路径

docx_NT = []
# 遍历文档中的段落
for i, duan in enumerate(doc.paragraphs):
    # 将每个段落的数字添加到列表中，不论段落是否为空
    docx_NT.append(i + 1)  # 因为Python的索引从0开始，所以加1

# 输出结果
print(docx_NT)

my_json_data = {
    'docx': f'{file}.docx',
    'ingT': [],
    'NT': docx_NT
}

while True: 
    # 检测列表的元素数是否超过100
    while True:
        if len(my_json_data['ingT']) >= xian_cheng_int:  # 检查线程是否大于100
            print(f"达到{xian_cheng_int}线程")
            print("当前线程数：",len(my_json_data['ingT']))
            time.sleep(1)
        else:
            print(f"未达到{xian_cheng_int}线程")
            print("当前线程数：",len(my_json_data['ingT']))
            break
    
    if my_json_data['NT'] == []:
        break

    with lock:
        j = min(my_json_data['NT'])  # 使用min函数找到最小值
        my_json_data['ingT'].append(j)  # 使用remove方法删除最小值
        my_json_data['NT'].remove(j)  # 使用remove方法删除最小值

    print("正在搞段落:",j)
    number_paragraph = j - 1
    # 检测是否有中英文
    paragraph = doc.paragraphs[number_paragraph]
    if paragraph.text.strip(): 
        print("段落包含文字")
        thread = threading.Thread(target=get,args=(j,number_paragraph,))  
        thread.start()
    else:
        print("段落不包含文字")
        with lock:
            my_json_data['ingT'].remove(j)

print("翻译全部完成,正在等待线程完成")

